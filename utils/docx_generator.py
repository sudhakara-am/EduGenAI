from docx import Document
from io import BytesIO


class DocxGenerator:

    @staticmethod
    def create_docx_buffer(
        title,
        content
    ):

        document = Document()

        document.add_heading(
            title,
            level=1
        )

        document.add_paragraph(
            content
        )

        buffer = BytesIO()

        document.save(
            buffer
        )

        buffer.seek(0)

        return buffer